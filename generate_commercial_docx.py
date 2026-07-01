import os
import subprocess
import sys

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Add Title
title = doc.add_heading('NEXUS OBSERVATORY', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Plataforma Integral de Observabilidad y Gestión para Modelos de Lenguaje (LLMs)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].bold = True

doc.add_paragraph('\n')

# Section 1
h1 = doc.add_heading('1. Resumen Ejecutivo y Propuesta de Valor', level=1)
doc.add_paragraph(
    'NEXUS Observatory es una solución empresarial de vanguardia diseñada para maximizar el retorno de inversión (ROI) '
    'y garantizar la seguridad operativa en aplicaciones impulsadas por Inteligencia Artificial y Modelos de Lenguaje '
    '(LLMs). En un entorno donde la IA transforma industrias, NEXUS otorga a los líderes tecnológicos (CTOs, líderes '
    'de ingeniería y directivos) una visibilidad absoluta sobre el rendimiento, los costos financieros y la calidad del '
    'código de sus ecosistemas.'
)
doc.add_paragraph(
    'Nuestra plataforma transforma la incertidumbre de los modelos generativos en métricas precisas y controlables, '
    'facilitando la toma de decisiones estratégicas, optimizando los recursos e impulsando la innovación de forma '
    'segura y escalable.'
)

# Section 2
doc.add_heading('2. Beneficios Principales y Casos de Uso Empresarial', level=1)
doc.add_paragraph(
    'NEXUS ha sido construido pensando en la eficiencia empresarial y la excelencia técnica, ofreciendo las siguientes '
    'soluciones clave:'
)

ul1 = doc.add_paragraph(style='List Bullet')
ul1.add_run('Control Financiero y Optimización de Costos (LLMOps): ').bold = True
ul1.add_run('Monitorización en tiempo real del consumo de tokens, latencias y costos operativos por cada interacción. Permite ajustar presupuestos y evitar sobrecostos en el uso de APIs comerciales.')

ul2 = doc.add_paragraph(style='List Bullet')
ul2.add_run('Auditoría Inteligente de Código: ').bold = True
ul2.add_run('Análisis automatizado impulsado por IA que evalúa la calidad, seguridad y mantenibilidad del código fuente de su organización. Detecta vulnerabilidades tempranas y asegura el cumplimiento de estándares internacionales.')

ul3 = doc.add_paragraph(style='List Bullet')
ul3.add_run('Interacción Contextual con Repositorios (Chat RAG): ').bold = True
ul3.add_run('Capacidad de dialogar directamente con el conocimiento técnico de la empresa. Facilita el "onboarding" de nuevos desarrolladores y democratiza el acceso a la información técnica.')

ul4 = doc.add_paragraph(style='List Bullet')
ul4.add_run('Pruebas Estratégicas (A/B Testing): ').bold = True
ul4.add_run('Plataforma empírica para evaluar diferentes modelos de inteligencia artificial y determinar de forma objetiva cuál ofrece el mejor balance entre costo, velocidad y precisión para su caso de uso específico.')

# Section 3
doc.add_heading('3. Arquitectura y Escalabilidad Tecnológica', level=1)
doc.add_paragraph(
    'El sistema Nexus Observatory se fundamenta en una arquitectura moderna, modular y altamente escalable, '
    'diseñado bajo el patrón cliente-servidor y optimizado para garantizar un rendimiento excepcional:'
)
doc.add_paragraph(
    '• Experiencia de Usuario Premium: Interfaz moderna e intuitiva que facilita flujos de trabajo eficientes y despliega dashboards analíticos interactivos de alta calidad visual.\n'
    '• Motor de Procesamiento Robusto: Un backend de alto rendimiento, capaz de gestionar peticiones concurrentes y asíncronas con mínima latencia.\n'
    '• Base de Conocimiento Inteligente: Integración de bases de datos especializadas para IA, permitiendo indexar el código fuente de la organización y ofrecer respuestas instantáneas y contextualizadas.'
)

# Section 4
doc.add_heading('4. Ventaja Competitiva y Seguridad Institucional', level=1)
doc.add_paragraph(
    'La adopción de NEXUS Observatory posiciona a su organización a la vanguardia tecnológica. Nuestra plataforma '
    'centraliza herramientas avanzadas, creando un entorno unificado donde los líderes pueden auditar el comportamiento '
    'algorítmico y financiero de forma transparente.'
)
doc.add_paragraph(
    'Además, su diseño prioriza la protección de la propiedad intelectual, incorporando mecanismos para garantizar '
    'que el análisis de código y las consultas corporativas se realicen bajo estrictos estándares de confidencialidad.'
)

# Section 5
doc.add_heading('5. Visión a Futuro y Roadmap Estratégico', level=1)
doc.add_paragraph(
    'NEXUS Observatory se encuentra en constante evolución para adaptarse a las exigentes demandas del mercado empresarial. '
    'Nuestro plan de desarrollo estratégico incluye:'
)
ul5 = doc.add_paragraph(style='List Bullet')
ul5.add_run('Despliegue Híbrido y Soberanía de Datos: ').bold = True
ul5.add_run('Integración con modelos de ejecución local para industrias altamente reguladas (como el sector financiero o de salud), garantizando privacidad total.')

ul6 = doc.add_paragraph(style='List Bullet')
ul6.add_run('Integración Nativa DevSecOps: ').bold = True
ul6.add_run('Acoplamiento automatizado con procesos de integración y entrega continua (CI/CD) para asegurar la calidad estructural del software antes de su paso a producción.')

ul7 = doc.add_paragraph(style='List Bullet')
ul7.add_run('Analítica Predictiva de Riesgos: ').bold = True
ul7.add_run('Módulos avanzados orientados a evaluar el historial de desarrollo para predecir, dimensionar y prevenir la deuda técnica a mediano y largo plazo.')

doc.save(r'C:\Users\cristian andres\OneDrive\Documentos\investigacion\Investigacion-1-Nexus-Observatory\Documentos Tecnicos\Especificacion_Tecnica_Nexus_Observatory.docx')
print("Document generated successfully.")
